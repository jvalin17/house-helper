"""Integration tests for DOCX format preservation through the full pipeline.

Tests: import stores binary → generate uses surgery → export returns preserved DOCX.
"""

import io
import json
import sqlite3
from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document
from docx.shared import Pt

from shared.db import connect_sync
from agents.job.repositories.knowledge_repo import KnowledgeRepository
from agents.job.services.knowledge import KnowledgeService
from agents.job.services.resume import ResumeService


def _make_test_docx(tmp_path: Path) -> Path:
    """Create a test DOCX with known formatting."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Jane Doe")
    run.font.name = "Garamond"
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph("jane@email.com")

    doc.add_heading("SUMMARY", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Backend engineer with 8 years experience.")
    run.font.name = "Garamond"
    run.font.size = Pt(11)

    doc.add_heading("WORK EXPERIENCE", level=2)
    p = doc.add_paragraph("TestCorp | Senior Engineer\tJan 2022 – Present")
    run = p.runs[0]
    run.font.name = "Garamond"
    run.bold = True

    for bullet in ["Built distributed systems", "Led team of 5 engineers"]:
        bp = doc.add_paragraph(style="List Bullet")
        brun = bp.add_run(bullet)
        brun.font.name = "Garamond"
        brun.font.size = Pt(10)

    doc.add_heading("EDUCATION", level=2)
    doc.add_paragraph("BS Computer Science, State University\t2018")

    path = tmp_path / "test_resume.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def db():
    conn = connect_sync(db_path=Path(":memory:"))
    return conn


@pytest.fixture
def knowledge_svc(db):
    repo = KnowledgeRepository(db)
    return KnowledgeService(repo, db)


class TestImportStoresDocxBinary:
    def test_docx_binary_stored_in_settings(self, db, knowledge_svc, tmp_path):
        docx_path = _make_test_docx(tmp_path)
        knowledge_svc.import_resume(docx_path)

        row = db.execute("SELECT value FROM settings WHERE key = 'original_resume_docx'").fetchone()
        assert row is not None

        import base64
        b64 = json.loads(row["value"])
        docx_bytes = base64.b64decode(b64)
        doc = Document(io.BytesIO(docx_bytes))
        assert len(doc.paragraphs) > 0

    def test_paragraph_map_stored_in_settings(self, db, knowledge_svc, tmp_path):
        docx_path = _make_test_docx(tmp_path)
        knowledge_svc.import_resume(docx_path)

        row = db.execute("SELECT value FROM settings WHERE key = 'original_resume_map'").fetchone()
        assert row is not None

        para_map = json.loads(row["value"])
        assert para_map["format"] == "docx"
        assert para_map["total_paragraphs"] > 0
        assert len(para_map["sections"]["experience"]["roles"]) >= 1

    def test_txt_import_does_not_store_docx(self, db, knowledge_svc, tmp_path):
        txt_path = tmp_path / "resume.txt"
        txt_path.write_text("WORK EXPERIENCE\nTestCorp | Engineer\tJan 2022 – Present\n- Built stuff\n")
        knowledge_svc.import_resume(txt_path)

        row = db.execute("SELECT value FROM settings WHERE key = 'original_resume_docx'").fetchone()
        assert row is None

    def test_raw_text_still_stored(self, db, knowledge_svc, tmp_path):
        docx_path = _make_test_docx(tmp_path)
        knowledge_svc.import_resume(docx_path)

        row = db.execute("SELECT value FROM settings WHERE key = 'original_resume'").fetchone()
        assert row is not None
        text = json.loads(row["value"])
        assert "Jane Doe" in text


class TestResumeExportWithDocxBinary:
    @staticmethod
    def _create_job(db):
        db.execute(
            "INSERT INTO jobs (id, title, company, parsed_data) VALUES (1, 'Test', 'TestCo', '{}')"
        )
        db.commit()

    def test_export_docx_returns_stored_binary(self, db, tmp_path):
        """If a resume has docx_binary, export('docx') returns it directly."""
        self._create_job(db)
        doc = Document()
        doc.add_paragraph("Test resume content")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        db.execute(
            "INSERT INTO resumes (job_id, content, preferences, docx_binary) VALUES (?, ?, ?, ?)",
            (1, "test content", "{}", docx_bytes),
        )
        db.commit()

        repo = KnowledgeRepository(db)
        svc = ResumeService(repo, db)
        result = svc.export(1, "docx")

        assert result == docx_bytes

    def test_export_md_ignores_docx_binary(self, db):
        """Markdown export uses text content, not DOCX binary."""
        self._create_job(db)
        db.execute(
            "INSERT INTO resumes (job_id, content, preferences, docx_binary) VALUES (?, ?, ?, ?)",
            (1, "# Test Resume\n\nContent here", "{}", b"fake docx bytes"),
        )
        db.commit()

        repo = KnowledgeRepository(db)
        svc = ResumeService(repo, db)
        result = svc.export(1, "md")

        assert b"Test Resume" in result
        assert b"fake docx bytes" not in result

    def test_export_docx_falls_back_without_binary(self, db):
        """If no docx_binary, falls back to markdown→DOCX conversion."""
        self._create_job(db)
        db.execute(
            "INSERT INTO resumes (job_id, content, preferences) VALUES (?, ?, ?)",
            (1, "# Test\n\n- Bullet point", "{}"),
        )
        db.commit()

        repo = KnowledgeRepository(db)
        svc = ResumeService(repo, db)
        result = svc.export(1, "docx")

        # Should be valid DOCX bytes
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0


class TestDocxSurgeryMissingData:
    """Tests for when DOCX binary or paragraph map are missing — must not crash."""

    @staticmethod
    def _create_job(db):
        db.execute("INSERT INTO jobs (id, title, company, parsed_data) VALUES (99, 'Test', 'Co', '{}')")
        db.commit()

    def test_generate_without_docx_stored(self, db):
        """Generate works when no DOCX was ever imported."""
        self._create_job(db)
        # Store only plain text resume, no DOCX binary
        db.execute(
            "INSERT OR REPLACE INTO settings (key, value) VALUES ('original_resume', ?)",
            (json.dumps("SUMMARY\nTest engineer\n\nWORK EXPERIENCE\nTestCo | Engineer\tJan 2020 – Present\n- Built stuff"),),
        )
        db.commit()

        repo = KnowledgeRepository(db)
        repo.save_experience(type="job", title="Engineer", company="TestCo",
                           start_date="2020-01", description="Built stuff")
        repo.save_skill(name="Python", category="languages")

        svc = ResumeService(repo, db)
        # No LLM, should use fallback builder
        result = svc.generate(job_id=99, preferences={})
        assert result["content"]
        assert result["id"]

    def test_has_original_docx_false_when_map_missing(self, db):
        """_has_original_docx returns False if map is missing even if binary exists."""
        db.execute(
            "INSERT OR REPLACE INTO settings (key, value) VALUES ('original_resume_docx', ?)",
            (json.dumps("fakebase64"),),
        )
        db.commit()

        repo = KnowledgeRepository(db)
        svc = ResumeService(repo, db)
        assert svc._has_original_docx() is False

    def test_has_original_docx_false_when_binary_missing(self, db):
        """_has_original_docx returns False if binary is missing even if map exists."""
        db.execute(
            "INSERT OR REPLACE INTO settings (key, value) VALUES ('original_resume_map', ?)",
            (json.dumps({"format": "docx", "total_paragraphs": 5, "sections": {}}),),
        )
        db.commit()

        repo = KnowledgeRepository(db)
        svc = ResumeService(repo, db)
        assert svc._has_original_docx() is False

    def test_get_paragraph_map_returns_none_when_missing(self, db):
        repo = KnowledgeRepository(db)
        svc = ResumeService(repo, db)
        assert svc._get_paragraph_map() is None

    def test_get_original_docx_returns_none_when_missing(self, db):
        repo = KnowledgeRepository(db)
        svc = ResumeService(repo, db)
        assert svc._get_original_docx() is None


class TestDeleteEndpoints:
    def test_delete_education(self, db):
        repo = KnowledgeRepository(db)
        edu_id = repo.save_education(institution="MIT", degree="BS", field="CS")
        assert len(repo.list_education()) == 1

        repo.delete_education(edu_id)
        assert len(repo.list_education()) == 0

    def test_delete_project(self, db):
        repo = KnowledgeRepository(db)
        proj_id = repo.save_project(name="TestProj", description="desc")
        assert len(repo.list_projects()) == 1

        repo.delete_project(proj_id)
        assert len(repo.list_projects()) == 0
