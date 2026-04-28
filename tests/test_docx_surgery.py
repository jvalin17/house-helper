"""Tests for DOCX surgical editing — paragraph map + format-preserving text replacement."""

import io

import pytest
from docx import Document
from docx.shared import Pt, RGBColor

from shared.docx_surgery import build_paragraph_map, apply_edits, _replace_text


# ── Fixtures ──────────────────────────────────


def _make_resume_doc() -> Document:
    """Create a realistic resume DOCX with known structure and formatting."""
    doc = Document()

    # Header / contact (before any section header)
    p = doc.add_paragraph()
    run = p.add_run("Jane Doe")
    run.font.name = "Garamond"
    run.font.size = Pt(16)
    run.bold = True

    p = doc.add_paragraph("jane@email.com | linkedin.com/janedoe | (555) 123-4567")

    # SUMMARY section
    h = doc.add_heading("SUMMARY", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Experienced backend engineer with 8 years building distributed systems.")
    run.font.name = "Garamond"
    run.font.size = Pt(11)

    # WORK EXPERIENCE section
    doc.add_heading("WORK EXPERIENCE", level=2)

    # Role 1: Zillow
    p = doc.add_paragraph("Zillow | Senior Software Engineer\tOct 2022 – Present")
    run = p.runs[0] if p.runs else p.add_run("")
    run.font.name = "Garamond"
    run.font.size = Pt(11)
    run.bold = True

    for bullet in [
        "Built notification pipeline processing 2M daily events",
        "Led migration to Kubernetes reducing deploy time by 60%",
        "Designed analytics dashboard used by 50+ stakeholders",
    ]:
        bp = doc.add_paragraph(style="List Bullet")
        brun = bp.add_run(bullet)
        brun.font.name = "Garamond"
        brun.font.size = Pt(10)

    # Role 2: Dematic
    p = doc.add_paragraph("Dematic | Software Engineer\tJan 2019 – Sep 2022")
    run = p.runs[0] if p.runs else p.add_run("")
    run.font.name = "Garamond"
    run.font.size = Pt(11)
    run.bold = True

    for bullet in [
        "Developed warehouse automation control system in Java",
        "Optimized path-finding algorithm reducing cycle time 25%",
    ]:
        bp = doc.add_paragraph(style="List Bullet")
        brun = bp.add_run(bullet)
        brun.font.name = "Garamond"
        brun.font.size = Pt(10)

    # EDUCATION section
    doc.add_heading("EDUCATION", level=2)
    doc.add_paragraph("BS Computer Science, Purdue University\t2018")

    # SKILLS section
    doc.add_heading("TECHNICAL SKILLS", level=2)
    doc.add_paragraph("Languages: Python, Java, Go, TypeScript")
    doc.add_paragraph("Infrastructure: Kubernetes, Docker, AWS, Terraform")

    return doc


def _doc_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def sample_doc():
    return _make_resume_doc()


@pytest.fixture
def sample_bytes():
    return _doc_to_bytes(_make_resume_doc())


@pytest.fixture
def sample_map(sample_doc):
    return build_paragraph_map(sample_doc)


# ── Paragraph Map Tests ──────────────────────


class TestBuildParagraphMap:
    def test_detects_summary_section(self, sample_map):
        summary = sample_map["sections"]["summary"]
        assert len(summary["paragraph_indices"]) >= 1

    def test_detects_experience_roles(self, sample_map):
        roles = sample_map["sections"]["experience"]["roles"]
        assert len(roles) == 2
        companies = [r["company"] for r in roles]
        assert "Zillow" in companies
        assert "Dematic" in companies

    def test_zillow_has_3_bullets(self, sample_map):
        roles = sample_map["sections"]["experience"]["roles"]
        zillow = next(r for r in roles if "Zillow" in r["company"])
        assert len(zillow["bullet_indices"]) == 3
        assert len(zillow["bullet_texts"]) == 3
        assert "notification pipeline" in zillow["bullet_texts"][0]

    def test_dematic_has_2_bullets(self, sample_map):
        roles = sample_map["sections"]["experience"]["roles"]
        dematic = next(r for r in roles if "Dematic" in r["company"])
        assert len(dematic["bullet_indices"]) == 2

    def test_detects_education(self, sample_map):
        edu = sample_map["sections"]["education"]
        assert len(edu["paragraph_indices"]) >= 1

    def test_detects_skills(self, sample_map):
        skills = sample_map["sections"]["skills"]
        assert len(skills["paragraph_indices"]) >= 1

    def test_total_paragraphs_matches(self, sample_doc, sample_map):
        assert sample_map["total_paragraphs"] == len(sample_doc.paragraphs)

    def test_format_is_docx(self, sample_map):
        assert sample_map["format"] == "docx"


# ── Text Replacement Tests ───────────────────


class TestReplaceText:
    def test_preserves_font_name(self):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("Original text")
        run.font.name = "Garamond"
        run.font.size = Pt(11)

        _replace_text(p, "New text")

        assert p.text == "New text"
        assert p.runs[0].font.name == "Garamond"
        assert p.runs[0].font.size == Pt(11)

    def test_preserves_bold(self):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("Bold text")
        run.bold = True

        _replace_text(p, "Still bold")

        assert p.text == "Still bold"
        assert p.runs[0].bold is True

    def test_preserves_color(self):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("Colored text")
        run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

        _replace_text(p, "New colored text")

        assert p.text == "New colored text"
        assert p.runs[0].font.color.rgb == RGBColor(0x2B, 0x57, 0x9A)

    def test_handles_empty_runs(self):
        doc = Document()
        p = doc.add_paragraph("No explicit runs")

        _replace_text(p, "Replaced")

        assert p.text == "Replaced"

    def test_clears_to_empty(self):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("Will be cleared")
        run.font.name = "Arial"

        _replace_text(p, "")

        assert p.text == ""


# ── Apply Edits Tests ────────────────────────


class TestApplyEdits:
    def test_replaces_summary(self, sample_bytes, sample_map):
        edits = {"summary": "New tailored summary for this role."}
        result = apply_edits(sample_bytes, sample_map, edits)
        assert result is not None

        doc = Document(io.BytesIO(result))
        summary_idx = sample_map["sections"]["summary"]["paragraph_indices"][0]
        assert doc.paragraphs[summary_idx].text == "New tailored summary for this role."

    def test_replaces_zillow_bullets(self, sample_bytes, sample_map):
        edits = {
            "experience_edits": [{
                "company": "Zillow",
                "title": "Senior Software Engineer",
                "bullets": [
                    "Architected event-driven notification system processing 3M daily events",
                    "Spearheaded Kubernetes migration cutting deploy time by 70%",
                    "Built real-time analytics dashboard for 100+ stakeholders",
                ],
                "swaps": [],
            }],
        }
        result = apply_edits(sample_bytes, sample_map, edits)
        assert result is not None

        doc = Document(io.BytesIO(result))
        roles = sample_map["sections"]["experience"]["roles"]
        zillow = next(r for r in roles if "Zillow" in r["company"])
        for j, idx in enumerate(zillow["bullet_indices"]):
            assert doc.paragraphs[idx].text == edits["experience_edits"][0]["bullets"][j]

    def test_preserves_font_after_edit(self, sample_bytes, sample_map):
        edits = {
            "experience_edits": [{
                "company": "Zillow",
                "title": "Senior Software Engineer",
                "bullets": ["New bullet with preserved formatting"],
                "swaps": [],
            }],
        }
        result = apply_edits(sample_bytes, sample_map, edits)
        doc = Document(io.BytesIO(result))

        roles = sample_map["sections"]["experience"]["roles"]
        zillow = next(r for r in roles if "Zillow" in r["company"])
        edited_para = doc.paragraphs[zillow["bullet_indices"][0]]
        assert edited_para.runs[0].font.name == "Garamond"

    def test_fewer_bullets_clears_excess(self, sample_bytes, sample_map):
        edits = {
            "experience_edits": [{
                "company": "Zillow",
                "title": "Senior Software Engineer",
                "bullets": ["Only one bullet now"],
                "swaps": [],
            }],
        }
        result = apply_edits(sample_bytes, sample_map, edits)
        doc = Document(io.BytesIO(result))

        roles = sample_map["sections"]["experience"]["roles"]
        zillow = next(r for r in roles if "Zillow" in r["company"])
        assert doc.paragraphs[zillow["bullet_indices"][0]].text == "Only one bullet now"
        assert doc.paragraphs[zillow["bullet_indices"][1]].text == ""
        assert doc.paragraphs[zillow["bullet_indices"][2]].text == ""

    def test_unmatched_company_untouched(self, sample_bytes, sample_map):
        edits = {
            "experience_edits": [{
                "company": "Google",
                "title": "SWE",
                "bullets": ["Google bullet"],
                "swaps": [],
            }],
        }
        result = apply_edits(sample_bytes, sample_map, edits)
        doc = Document(io.BytesIO(result))

        roles = sample_map["sections"]["experience"]["roles"]
        zillow = next(r for r in roles if "Zillow" in r["company"])
        assert "notification pipeline" in doc.paragraphs[zillow["bullet_indices"][0]].text

    def test_integrity_check_returns_none(self, sample_bytes, sample_map):
        bad_map = {**sample_map, "total_paragraphs": 999}
        result = apply_edits(sample_bytes, bad_map, {})
        assert result is None

    def test_strips_bullet_prefix(self, sample_bytes, sample_map):
        edits = {
            "experience_edits": [{
                "company": "Zillow",
                "title": "Senior Software Engineer",
                "bullets": ["- Bullet with dash prefix", "• Bullet with dot prefix", "No prefix"],
                "swaps": [],
            }],
        }
        result = apply_edits(sample_bytes, sample_map, edits)
        doc = Document(io.BytesIO(result))

        roles = sample_map["sections"]["experience"]["roles"]
        zillow = next(r for r in roles if "Zillow" in r["company"])
        assert doc.paragraphs[zillow["bullet_indices"][0]].text == "Bullet with dash prefix"
        assert doc.paragraphs[zillow["bullet_indices"][1]].text == "Bullet with dot prefix"
        assert doc.paragraphs[zillow["bullet_indices"][2]].text == "No prefix"

    def test_empty_edits_returns_unchanged(self, sample_bytes, sample_map):
        result = apply_edits(sample_bytes, sample_map, {})
        assert result is not None
        # Should be valid DOCX
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) == sample_map["total_paragraphs"]

    def test_dematic_edit_doesnt_affect_zillow(self, sample_bytes, sample_map):
        edits = {
            "experience_edits": [{
                "company": "Dematic",
                "title": "Software Engineer",
                "bullets": ["Changed Dematic bullet"],
                "swaps": [],
            }],
        }
        result = apply_edits(sample_bytes, sample_map, edits)
        doc = Document(io.BytesIO(result))

        roles = sample_map["sections"]["experience"]["roles"]
        zillow = next(r for r in roles if "Zillow" in r["company"])
        # Zillow should be untouched
        assert "notification pipeline" in doc.paragraphs[zillow["bullet_indices"][0]].text
