"""TDD tests for multi-resume template system.

Users can store up to 5 resume files as templates.
One template is marked as default for matching/generation.
Import creates a template entry alongside KB data.
"""

import io
import json
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from shared.db import connect_sync
from agents.job.repositories.knowledge_repo import KnowledgeRepository


@pytest.fixture
def db():
    return connect_sync(db_path=Path(":memory:"))


def _make_test_docx(name: str = "test_resume.docx") -> tuple[bytes, str]:
    """Create a minimal test DOCX and return (bytes, filename)."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Test Name")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    doc.add_heading("SUMMARY", level=2)
    doc.add_paragraph("Test engineer summary")
    doc.add_heading("WORK EXPERIENCE", level=2)
    doc.add_paragraph("TestCo | Engineer\tJan 2020 – Present")
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run("Built stuff")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), name


class TestTemplateRepo:
    def test_save_template(self, db):
        from agents.job.repositories.template_repo import ResumeTemplateRepo

        repo = ResumeTemplateRepo(db)
        docx_bytes, _ = _make_test_docx()
        template_id = repo.save_template(
            name="Backend Resume",
            filename="backend_resume.docx",
            format="docx",
            raw_text="Test resume text",
            docx_binary=docx_bytes,
            paragraph_map={"format": "docx", "total_paragraphs": 5},
        )
        assert template_id > 0

    def test_list_templates(self, db):
        from agents.job.repositories.template_repo import ResumeTemplateRepo

        repo = ResumeTemplateRepo(db)
        docx_bytes, _ = _make_test_docx()
        repo.save_template("Resume A", "a.docx", "docx", "text a", docx_bytes)
        repo.save_template("Resume B", "b.docx", "docx", "text b", docx_bytes)

        templates = repo.list_templates()
        assert len(templates) == 2
        # Should not include docx_binary in list (too large)
        assert "docx_binary" not in templates[0]

    def test_get_template_with_binary(self, db):
        from agents.job.repositories.template_repo import ResumeTemplateRepo

        repo = ResumeTemplateRepo(db)
        docx_bytes, _ = _make_test_docx()
        tid = repo.save_template("Resume", "r.docx", "docx", "text", docx_bytes)

        template = repo.get_template(tid)
        assert template is not None
        assert template["docx_binary"] == docx_bytes

    def test_set_default(self, db):
        from agents.job.repositories.template_repo import ResumeTemplateRepo

        repo = ResumeTemplateRepo(db)
        docx_bytes, _ = _make_test_docx()
        t1 = repo.save_template("Resume A", "a.docx", "docx", "text a", docx_bytes)
        t2 = repo.save_template("Resume B", "b.docx", "docx", "text b", docx_bytes)

        repo.set_default(t2)

        templates = repo.list_templates()
        default = [t for t in templates if t["is_default"]]
        assert len(default) == 1
        assert default[0]["id"] == t2

    def test_delete_template(self, db):
        from agents.job.repositories.template_repo import ResumeTemplateRepo

        repo = ResumeTemplateRepo(db)
        docx_bytes, _ = _make_test_docx()
        tid = repo.save_template("Resume", "r.docx", "docx", "text", docx_bytes)

        repo.delete_template(tid)
        assert repo.list_templates() == []

    def test_max_5_templates(self, db):
        from agents.job.repositories.template_repo import ResumeTemplateRepo

        repo = ResumeTemplateRepo(db)
        docx_bytes, _ = _make_test_docx()
        for i in range(5):
            repo.save_template(f"Resume {i}", f"r{i}.docx", "docx", f"text {i}", docx_bytes)

        assert len(repo.list_templates()) == 5

        # 6th should raise
        with pytest.raises(ValueError, match="maximum"):
            repo.save_template("Resume 6", "r6.docx", "docx", "text 6", docx_bytes)

    def test_get_default_template(self, db):
        from agents.job.repositories.template_repo import ResumeTemplateRepo

        repo = ResumeTemplateRepo(db)
        docx_bytes, _ = _make_test_docx()
        t1 = repo.save_template("Resume A", "a.docx", "docx", "text a", docx_bytes)
        repo.set_default(t1)

        default = repo.get_default_template()
        assert default is not None
        assert default["id"] == t1

    def test_no_default_returns_none(self, db):
        from agents.job.repositories.template_repo import ResumeTemplateRepo

        repo = ResumeTemplateRepo(db)
        assert repo.get_default_template() is None

    def test_first_template_auto_default(self, db):
        """First template uploaded should automatically be the default."""
        from agents.job.repositories.template_repo import ResumeTemplateRepo

        repo = ResumeTemplateRepo(db)
        docx_bytes, _ = _make_test_docx()
        repo.save_template("First Resume", "first.docx", "docx", "text", docx_bytes)

        templates = repo.list_templates()
        assert templates[0]["is_default"] == 1
