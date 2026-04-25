"""Integration tests for the FastAPI application — end-to-end API tests."""

import sqlite3

import pytest
from fastapi import FastAPI
from fastapi.testclient import TestClient

from shared.db import migrate
from coordinator import Coordinator


@pytest.fixture
def app(tmp_path):
    """Create a test FastAPI app with an in-memory-like DB."""
    db_path = tmp_path / "test.db"
    conn = sqlite3.connect(str(db_path), check_same_thread=False)
    conn.row_factory = sqlite3.Row
    migrate(conn)

    test_app = FastAPI()
    coordinator = Coordinator(conn=conn, llm_provider=None)
    test_app.include_router(coordinator.get_router())

    @test_app.get("/health")
    def health():
        return {"status": "ok"}

    yield test_app
    conn.close()


@pytest.fixture
def client(app):
    return TestClient(app)


class TestHealth:
    def test_health_check(self, client):
        r = client.get("/health")
        assert r.status_code == 200
        assert r.json()["status"] == "ok"


class TestKnowledgeBank:
    def test_create_and_list_entries(self, client):
        r = client.post("/api/knowledge/entries", json={
            "type": "job", "title": "Engineer", "company": "Acme",
            "description": "Built APIs",
        })
        assert r.status_code == 200
        assert r.json()["id"] > 0

        r = client.get("/api/knowledge/entries")
        assert r.status_code == 200
        assert len(r.json()["experiences"]) == 1

    def test_update_entry(self, client):
        r = client.post("/api/knowledge/entries", json={
            "type": "job", "title": "Old", "company": "Co",
        })
        entry_id = r.json()["id"]

        r = client.put(f"/api/knowledge/entries/{entry_id}", json={"title": "New"})
        assert r.status_code == 200
        assert r.json()["title"] == "New"

    def test_delete_entry(self, client):
        r = client.post("/api/knowledge/entries", json={
            "type": "job", "title": "Del", "company": "Co",
        })
        entry_id = r.json()["id"]
        r = client.delete(f"/api/knowledge/entries/{entry_id}")
        assert r.status_code == 200

    def test_create_and_list_skills(self, client):
        client.post("/api/knowledge/skills", json={"name": "Python", "category": "language"})
        r = client.get("/api/knowledge/skills")
        assert r.status_code == 200
        assert len(r.json()) == 1

    def test_extract_skills_from_text(self, client):
        r = client.post("/api/knowledge/extract", json={
            "text": "Experience with Python, React, and Docker",
        })
        assert r.status_code == 200
        skills = r.json()["extracted_skills"]
        assert "Python" in skills


class TestResumeImport:
    def test_upload_docx_resume(self, client):
        import io
        from docx import Document

        doc = Document()
        doc.add_paragraph("John Doe").runs[0].bold = True
        doc.add_paragraph("WORK EXPERIENCE").runs[0].bold = True
        p = doc.add_paragraph("Acme | Engineer\tJan 2020 – Dec 2023")
        p.runs[0].bold = True

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        r = client.post(
            "/api/knowledge/import",
            files={"file": ("resume.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert r.status_code == 200
        data = r.json()
        assert data["experiences"] >= 1

    def test_upload_txt_resume(self, client):
        content = b"WORK EXPERIENCE\nAcme | Engineer\tJan 2020 - Dec 2023\n- Built APIs"
        r = client.post(
            "/api/knowledge/import",
            files={"file": ("resume.txt", content, "text/plain")},
        )
        assert r.status_code == 200

    def test_upload_unsupported_format(self, client):
        r = client.post(
            "/api/knowledge/import",
            files={"file": ("resume.jpg", b"fake image", "image/jpeg")},
        )
        assert r.status_code == 400


class TestJobs:
    def test_parse_and_list_jobs(self, client):
        r = client.post("/api/jobs/parse", json={
            "inputs": ["Software Engineer at BigTech\nRequirements:\n- Python\n- React"],
        })
        assert r.status_code == 200
        assert len(r.json()["jobs"]) == 1

        r = client.get("/api/jobs")
        assert len(r.json()) == 1

    def test_get_and_delete_job(self, client):
        r = client.post("/api/jobs/parse", json={
            "inputs": ["Engineer at Co\n- Java"],
        })
        job_id = r.json()["jobs"][0]["id"]

        r = client.get(f"/api/jobs/{job_id}")
        assert r.status_code == 200

        r = client.delete(f"/api/jobs/{job_id}")
        assert r.status_code == 200

        r = client.get(f"/api/jobs/{job_id}")
        assert r.status_code == 404


class TestMatching:
    def _setup_knowledge(self, client):
        client.post("/api/knowledge/entries", json={
            "type": "job", "title": "Eng", "company": "X", "description": "Python dev",
        })
        client.post("/api/knowledge/skills", json={"name": "Python", "category": "language"})
        client.post("/api/knowledge/skills", json={"name": "React", "category": "framework"})

    def test_match_single_job(self, client):
        self._setup_knowledge(client)
        r = client.post("/api/jobs/parse", json={
            "inputs": ["SWE at Co\n- Python\n- React"],
        })
        job_id = r.json()["jobs"][0]["id"]

        r = client.post(f"/api/jobs/{job_id}/match")
        assert r.status_code == 200
        assert "score" in r.json()

    def test_match_batch(self, client):
        self._setup_knowledge(client)
        r = client.post("/api/jobs/parse", json={
            "inputs": ["SWE at A\n- Python", "Chef at B\n- Cooking"],
        })
        ids = [j["id"] for j in r.json()["jobs"]]

        r = client.post("/api/jobs/match-batch", json={"job_ids": ids})
        assert r.status_code == 200
        assert len(r.json()["results"]) == 2


class TestResumes:
    def _setup(self, client):
        client.post("/api/knowledge/entries", json={
            "type": "job", "title": "Eng", "company": "Acme", "description": "Built APIs",
        })
        client.post("/api/knowledge/skills", json={"name": "Python", "category": "language"})
        r = client.post("/api/jobs/parse", json={"inputs": ["SWE at Co\n- Python"]})
        return r.json()["jobs"][0]["id"]

    def test_generate_resume(self, client):
        job_id = self._setup(client)
        r = client.post("/api/resumes/generate", json={"job_id": job_id})
        assert r.status_code == 200
        assert "Acme" in r.json()["content"]

    def test_export_resume_pdf(self, client):
        job_id = self._setup(client)
        r = client.post("/api/resumes/generate", json={"job_id": job_id})
        resume_id = r.json()["id"]

        r = client.get(f"/api/resumes/{resume_id}/export?format=pdf")
        assert r.status_code == 200
        assert r.content[:5] == b"%PDF-"

    def test_resume_feedback(self, client):
        job_id = self._setup(client)
        r = client.post("/api/resumes/generate", json={"job_id": job_id})
        resume_id = r.json()["id"]

        r = client.post(f"/api/resumes/{resume_id}/feedback", json={"rating": 1})
        assert r.status_code == 200


class TestCoverLetters:
    def _setup(self, client):
        client.post("/api/knowledge/entries", json={
            "type": "job", "title": "Eng", "company": "Acme", "description": "Built APIs",
        })
        client.post("/api/knowledge/skills", json={"name": "Python", "category": "language"})
        r = client.post("/api/jobs/parse", json={"inputs": ["SWE at BigTech\n- Python"]})
        return r.json()["jobs"][0]["id"]

    def test_generate_cover_letter(self, client):
        job_id = self._setup(client)
        r = client.post("/api/cover-letters/generate", json={"job_id": job_id})
        assert r.status_code == 200
        assert "BigTech" in r.json()["content"]

    def test_edit_cover_letter(self, client):
        job_id = self._setup(client)
        r = client.post("/api/cover-letters/generate", json={"job_id": job_id})
        cl_id = r.json()["id"]

        r = client.put(f"/api/cover-letters/{cl_id}", json={"content": "Edited content"})
        assert r.status_code == 200
        assert r.json()["content"] == "Edited content"

    def test_export_cover_letter_docx(self, client):
        job_id = self._setup(client)
        r = client.post("/api/cover-letters/generate", json={"job_id": job_id})
        cl_id = r.json()["id"]

        r = client.get(f"/api/cover-letters/{cl_id}/export?format=docx")
        assert r.status_code == 200
        assert r.content[:2] == b"PK"


class TestApplications:
    def _setup(self, client):
        r = client.post("/api/jobs/parse", json={"inputs": ["SWE at Co\n- Python"]})
        return r.json()["jobs"][0]["id"]

    def test_create_and_list(self, client):
        job_id = self._setup(client)
        r = client.post("/api/applications", json={"job_id": job_id})
        assert r.status_code == 200
        assert r.json()["status"] == "applied"

        r = client.get("/api/applications")
        assert len(r.json()) == 1

    def test_update_status(self, client):
        job_id = self._setup(client)
        r = client.post("/api/applications", json={"job_id": job_id})
        app_id = r.json()["id"]

        r = client.put(f"/api/applications/{app_id}", json={"status": "interview"})
        assert r.json()["status"] == "interview"

    def test_status_history(self, client):
        job_id = self._setup(client)
        r = client.post("/api/applications", json={"job_id": job_id})
        app_id = r.json()["id"]
        client.put(f"/api/applications/{app_id}", json={"status": "interview"})
        client.put(f"/api/applications/{app_id}", json={"status": "offer"})

        r = client.get(f"/api/applications/{app_id}/history")
        statuses = [h["status"] for h in r.json()]
        assert "applied" in statuses
        assert "interview" in statuses
        assert "offer" in statuses


class TestCalibration:
    def _setup(self, client):
        client.post("/api/knowledge/skills", json={"name": "Python", "category": "language"})
        r = client.post("/api/jobs/parse", json={"inputs": ["SWE at Co\n- Python"]})
        job_id = r.json()["jobs"][0]["id"]
        client.post(f"/api/jobs/{job_id}/match")
        return job_id

    def test_submit_judgement(self, client):
        job_id = self._setup(client)
        r = client.post("/api/calibration/judge", json={
            "job_id": job_id, "rating": "good",
        })
        assert r.status_code == 200

    def test_recalculate_weights(self, client):
        job_id = self._setup(client)
        client.post("/api/calibration/judge", json={"job_id": job_id, "rating": "good"})
        r = client.post("/api/calibration/recalculate")
        assert r.status_code == 200
        weights = r.json()
        assert abs(sum(weights.values()) - 1.0) < 0.01


class TestPreferences:
    def test_set_and_get(self, client):
        r = client.put("/api/preferences", json={"tone": "formal", "length": "1_page"})
        assert r.status_code == 200

        r = client.get("/api/preferences")
        assert r.json()["tone"] == "formal"
