"""Shared fixtures for end-to-end workflow integration tests.

Each test gets a fresh FastAPI app wired to its own SQLite DB so workflows
can be exercised through the real HTTP surface (TestClient) without touching
external services. A deterministic mock LLM provider lets us assert against
features that require AI without spending real tokens.
"""

from __future__ import annotations

import io
import json
import sqlite3
from pathlib import Path
from typing import Callable

import pytest
from fastapi import FastAPI
from fastapi.testclient import TestClient

from shared.db import migrate
from shared.job_boards.base import JobBoardPlugin, JobResult, SearchFilters
from coordinator import Coordinator


class MockLLMProvider:
    """Deterministic in-memory LLM stand-in.

    Routes responses by the `feature` kwarg the production code passes (e.g.
    "resume_analyze", "resume_gen"). Tests can override responses per-feature
    to drive specific code paths.
    """

    def __init__(self, responses: dict[str, str] | None = None) -> None:
        self._responses = responses or {}
        self.calls: list[dict] = []

    def set_response(self, feature: str, response: str) -> None:
        self._responses[feature] = response

    def is_configured(self) -> bool:
        return True

    def provider_name(self) -> str:
        return "mock"

    def model_name(self) -> str:
        return "mock-model"

    def complete(self, prompt: str, system: str | None = None, feature: str = "unknown") -> str:
        self.calls.append({"prompt": prompt, "system": system, "feature": feature})
        if feature in self._responses:
            return self._responses[feature]
        return self._default_response(feature)

    @staticmethod
    def _default_response(feature: str) -> str:
        if feature == "resume_analyze":
            return json.dumps({
                "current_resume_match": 0.62,
                "knowledge_bank_match": 0.81,
                "match_gap": "moderate",
                "strengths": ["Python", "APIs"],
                "gaps": ["Kubernetes"],
                "suggested_improvements": [
                    {
                        "type": "bullet_rewrite",
                        "description": "Quantify Acme API impact",
                        "impact": "shows scale",
                        "source": "experience",
                    }
                ],
                "summary": "Solid match — quantify impact for stronger fit.",
            })
        if feature == "skill_extract":
            return json.dumps(["Python", "FastAPI", "React"])
        if feature == "resume_gen":
            return (
                "# Generated Resume\n\n"
                "## Summary\nSenior engineer with API expertise.\n\n"
                "## Experience\n- Acme: Built APIs in Python.\n"
            )
        if feature == "cover_letter":
            return "Dear Hiring Manager,\n\nI am excited to apply...\n\nSincerely,\nCandidate"
        return "{}"


class _MockJobBoard:
    """Minimal JobBoardPlugin that yields canned results without HTTP."""

    def __init__(self, results: list[JobResult]):
        self._results = results

    def search(self, filters: SearchFilters) -> list[JobResult]:
        return list(self._results)

    def is_available(self) -> bool:
        return True

    def requires_api_key(self) -> bool:
        return False

    def board_name(self) -> str:
        return "mock"


def _build_app(conn: sqlite3.Connection, llm_provider) -> FastAPI:
    app = FastAPI()
    coordinator = Coordinator(conn=conn, llm_provider=llm_provider)
    app.include_router(coordinator.get_router())
    return app


@pytest.fixture
def db(tmp_path):
    """Per-test SQLite database with full migrations applied."""
    conn = sqlite3.connect(str(tmp_path / "wf.db"), check_same_thread=False)
    conn.row_factory = sqlite3.Row
    migrate(conn)
    yield conn
    conn.close()


@pytest.fixture
def llm() -> MockLLMProvider:
    return MockLLMProvider()


@pytest.fixture
def app(db, llm) -> FastAPI:
    return _build_app(db, llm)


@pytest.fixture
def app_no_llm(db) -> FastAPI:
    """App without an LLM provider — exercises algorithmic fallbacks."""
    return _build_app(db, None)


@pytest.fixture
def client(app) -> TestClient:
    return TestClient(app)


@pytest.fixture
def client_no_llm(app_no_llm) -> TestClient:
    return TestClient(app_no_llm)


@pytest.fixture
def mock_job_boards(monkeypatch):
    """Patch the job-board factory to yield canned results.

    Returns a setter so each test can configure the JobResult list it wants.
    """

    state: dict = {"results": []}

    def _set_results(results: list[JobResult]) -> None:
        state["results"] = list(results)

    def _fake_get_available_boards():
        return [_MockJobBoard(state["results"])]

    monkeypatch.setattr(
        "shared.job_boards.factory.get_available_boards",
        _fake_get_available_boards,
    )
    monkeypatch.setattr(
        "agents.job.services.auto_search.get_available_boards",
        _fake_get_available_boards,
    )
    return _set_results


# ── DOCX / TXT resume helpers ──────────────────────────────────────────────


def _build_docx_resume(
    contact: str = "Jane Doe",
    experiences: list[tuple[str, str, str, list[str]]] | None = None,
    skills: list[str] | None = None,
    education: list[tuple[str, str]] | None = None,
    projects: list[tuple[str, str]] | None = None,
) -> bytes:
    """Construct a minimal DOCX byte string parseable by knowledge import."""
    from docx import Document

    doc = Document()
    p = doc.add_paragraph(contact)
    p.runs[0].bold = True

    exps = experiences or [("Acme Corp", "Senior Engineer", "Jan 2020 – Present",
                            ["Built REST APIs in Python", "Mentored 4 engineers"])]
    skills = skills or ["Python", "FastAPI", "React", "PostgreSQL"]
    education = education or [("State University", "BS Computer Science")]
    projects = projects or [("Resume Helper", "Resume tailor for job seekers")]

    head = doc.add_paragraph("WORK EXPERIENCE")
    head.runs[0].bold = True
    for company, title, dates, bullets in exps:
        line = doc.add_paragraph(f"{company} | {title}\t{dates}")
        line.runs[0].bold = True
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    head = doc.add_paragraph("SKILLS")
    head.runs[0].bold = True
    doc.add_paragraph(", ".join(skills))

    head = doc.add_paragraph("EDUCATION")
    head.runs[0].bold = True
    for institution, degree in education:
        doc.add_paragraph(f"{institution} — {degree}")

    head = doc.add_paragraph("PROJECTS")
    head.runs[0].bold = True
    for name, desc in projects:
        doc.add_paragraph(f"{name}: {desc}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def docx_resume_bytes() -> Callable[..., bytes]:
    return _build_docx_resume


@pytest.fixture
def txt_resume_bytes() -> bytes:
    return (
        "WORK EXPERIENCE\n"
        "Acme Corp | Senior Engineer\tJan 2020 - Present\n"
        "- Built REST APIs in Python\n"
        "- Mentored 4 engineers\n"
        "SKILLS\n"
        "Python, FastAPI, React, PostgreSQL\n"
        "EDUCATION\n"
        "State University — BS Computer Science\n"
    ).encode("utf-8")


@pytest.fixture
def seeded_client(client) -> TestClient:
    """Client with a baseline knowledge bank + a parsed job in the DB."""
    client.post("/api/knowledge/entries", json={
        "type": "job", "title": "Senior Engineer", "company": "Acme Corp",
        "start_date": "2020-01-01", "description": "Built APIs in Python and React",
    })
    for skill in ("Python", "FastAPI", "React"):
        client.post("/api/knowledge/skills", json={"name": skill, "category": "language"})
    client.post("/api/jobs/parse", json={
        "inputs": ["Backend Engineer at BigTech\nRequirements:\n- Python\n- React\n- FastAPI"],
    })
    return client


def make_job_result(
    title: str = "Backend Engineer",
    company: str = "BigTech",
    description: str = "Python, FastAPI, REST APIs",
    location: str = "Remote",
    url: str | None = None,
    salary: str | None = None,
    source: str = "mock",
) -> JobResult:
    """Helper for constructing JobResult instances from tests."""
    return JobResult(
        title=title,
        company=company,
        description=description,
        location=location,
        url=url or f"https://jobs.example.com/{title.lower().replace(' ', '-')}",
        salary=salary,
        source=source,
    )


def fixture_path() -> Path:
    """Helper returning the integration tests directory for fixtures."""
    return Path(__file__).resolve().parent
