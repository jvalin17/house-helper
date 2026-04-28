"""TDD tests for template integration with generate/export flow.

Core behavior:
1. Importing a resume creates a template entry
2. Importing a second resume creates a second template (not overwrite)
3. generate() uses the default template's text, not settings key
4. export(docx) uses the default template's DOCX binary
5. User can switch default template and generate uses the new one
"""

import io
import json
import tempfile
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest
from docx import Document
from docx.shared import Pt

from shared.db import connect_sync
from agents.job.repositories.knowledge_repo import KnowledgeRepository
from agents.job.repositories.template_repo import ResumeTemplateRepo
from agents.job.services.knowledge import KnowledgeService
from agents.job.services.resume import ResumeService


@pytest.fixture
def db():
    return connect_sync(db_path=Path(":memory:"))


@pytest.fixture
def template_repo(db):
    return ResumeTemplateRepo(db)


def _make_docx(name: str, summary: str) -> Path:
    """Create a test DOCX with identifiable content."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.font.name = "Arial"
    run.font.size = Pt(14)
    doc.add_heading("SUMMARY", level=2)
    p = doc.add_paragraph()
    run = p.add_run(summary)
    run.font.name = "Arial"
    doc.add_heading("WORK EXPERIENCE", level=2)
    doc.add_paragraph(f"TestCo | Engineer\tJan 2020 – Present")
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run("Built stuff at TestCo")

    tmp = Path(tempfile.mktemp(suffix=".docx"))
    doc.save(str(tmp))
    return tmp


class TestImportCreatesTemplate:
    def test_first_import_creates_template(self, db, template_repo):
        repo = KnowledgeRepository(db)
        svc = KnowledgeService(repo, db)
        docx_path = _make_docx("Jane Doe", "Backend engineer")
        try:
            svc.import_resume(docx_path)
        finally:
            docx_path.unlink(missing_ok=True)

        templates = template_repo.list_templates()
        assert len(templates) == 1
        assert templates[0]["is_default"] == 1

    def test_second_import_creates_second_template(self, db, template_repo):
        repo = KnowledgeRepository(db)
        svc = KnowledgeService(repo, db)

        path1 = _make_docx("Jane Doe", "Backend engineer v1")
        path2 = _make_docx("Jane Doe", "Backend engineer v2")
        try:
            svc.import_resume(path1)
            svc.import_resume(path2)
        finally:
            path1.unlink(missing_ok=True)
            path2.unlink(missing_ok=True)

        templates = template_repo.list_templates()
        assert len(templates) == 2

    def test_second_import_becomes_default(self, db, template_repo):
        repo = KnowledgeRepository(db)
        svc = KnowledgeService(repo, db)

        path1 = _make_docx("Jane v1", "Version 1")
        path2 = _make_docx("Jane v2", "Version 2")
        try:
            svc.import_resume(path1)
            svc.import_resume(path2)
        finally:
            path1.unlink(missing_ok=True)
            path2.unlink(missing_ok=True)

        default = template_repo.get_default_template()
        assert default is not None
        assert "version 2" in (default.get("raw_text") or "").lower()


class TestGenerateUsesTemplate:
    def test_generate_uses_default_template_text(self, db, template_repo):
        """generate() should use the default template's raw_text, not settings key."""
        repo = KnowledgeRepository(db)
        repo.save_experience(type="job", title="Engineer", company="TestCo",
                           start_date="2020-01", description="Built stuff")
        repo.save_skill(name="Python", category="languages")

        # Create a job
        db.execute("INSERT INTO jobs (id, title, company, parsed_data) VALUES (1, 'SWE', 'Acme', '{}')")
        db.commit()

        # Store template with known text
        docx_bytes, _ = io.BytesIO(), None
        doc = Document()
        doc.add_paragraph("TEMPLATE RESUME TEXT")
        buf = io.BytesIO()
        doc.save(buf)

        template_repo.save_template(
            name="My Resume",
            filename="resume.docx",
            file_format="docx",
            raw_text="SUMMARY\nTemplate resume for testing\n\nWORK EXPERIENCE\nTestCo | Engineer\tJan 2020\n- Built stuff",
            docx_binary=buf.getvalue(),
        )

        # Also store DIFFERENT text in settings (old path)
        db.execute(
            "INSERT OR REPLACE INTO settings (key, value) VALUES ('original_resume', ?)",
            (json.dumps("OLD SETTINGS RESUME - should NOT be used"),),
        )
        db.commit()

        svc = ResumeService(repo, db)
        # Without LLM, uses fallback builder — but _get_original_resume should return template text
        original = svc._get_original_resume()
        assert "Template resume" in original
        assert "OLD SETTINGS" not in original


class TestSwitchDefaultTemplate:
    def test_switch_default_changes_generate_text(self, db, template_repo):
        """Switching default template should change what _get_original_resume returns."""
        repo = KnowledgeRepository(db)
        svc = ResumeService(repo, db)

        # Create two templates
        doc = Document()
        doc.add_paragraph("dummy")
        buf = io.BytesIO()
        doc.save(buf)
        dummy_bytes = buf.getvalue()

        t1 = template_repo.save_template("Resume V1", "v1.docx", file_format="docx",
                                         raw_text="VERSION ONE resume text", docx_binary=dummy_bytes)
        t2 = template_repo.save_template("Resume V2", "v2.docx", file_format="docx",
                                         raw_text="VERSION TWO resume text", docx_binary=dummy_bytes)

        # t1 is auto-default (first)
        svc._default_template_cache = None
        assert "VERSION ONE" in svc._get_original_resume()

        # Switch to t2
        template_repo.set_default(t2)
        svc._default_template_cache = None
        assert "VERSION TWO" in svc._get_original_resume()

        # Switch back to t1
        template_repo.set_default(t1)
        svc._default_template_cache = None
        assert "VERSION ONE" in svc._get_original_resume()


class TestFallbackToSettings:
    def test_no_template_falls_back_to_settings(self, db):
        """When no templates exist, _get_original_resume uses settings key."""
        repo = KnowledgeRepository(db)
        svc = ResumeService(repo, db)

        # Store text in settings only (no templates)
        db.execute(
            "INSERT OR REPLACE INTO settings (key, value) VALUES ('original_resume', ?)",
            (json.dumps("SETTINGS FALLBACK TEXT"),),
        )
        db.commit()

        svc._default_template_cache = None
        result = svc._get_original_resume()
        assert result == "SETTINGS FALLBACK TEXT"

    def test_template_preferred_over_settings(self, db):
        """When template exists, it takes priority over settings key."""
        repo = KnowledgeRepository(db)
        template_repo = ResumeTemplateRepo(db)
        svc = ResumeService(repo, db)

        # Store in settings
        db.execute(
            "INSERT OR REPLACE INTO settings (key, value) VALUES ('original_resume', ?)",
            (json.dumps("OLD SETTINGS TEXT"),),
        )
        db.commit()

        # Create template
        doc = Document()
        doc.add_paragraph("dummy")
        buf = io.BytesIO()
        doc.save(buf)
        template_repo.save_template("Current", "current.docx", file_format="docx",
                                    raw_text="TEMPLATE TEXT", docx_binary=buf.getvalue())

        svc._default_template_cache = None
        result = svc._get_original_resume()
        assert result == "TEMPLATE TEXT"
        assert "OLD SETTINGS" not in result


class TestExportUsesTemplate:
    def test_export_docx_uses_default_template_binary(self, db, template_repo):
        """export(docx) should use the default template's DOCX binary."""
        repo = KnowledgeRepository(db)

        # Create a template with known DOCX
        doc = Document()
        doc.add_paragraph("TEMPLATE DOCX CONTENT")
        buf = io.BytesIO()
        doc.save(buf)
        template_binary = buf.getvalue()

        template_repo.save_template(
            name="My Resume",
            filename="resume.docx",
            file_format="docx",
            raw_text="template text",
            docx_binary=template_binary,
            paragraph_map={"format": "docx", "total_paragraphs": 1, "sections": {}},
        )

        # Create job + resume with docx_binary from surgery
        db.execute("INSERT INTO jobs (id, title, company, parsed_data) VALUES (1, 'SWE', 'Acme', '{}')")
        db.execute(
            "INSERT INTO resumes (id, job_id, content, preferences, docx_binary) VALUES (1, 1, 'content', '{}', ?)",
            (template_binary,),
        )
        db.commit()

        svc = ResumeService(repo, db)
        result = svc.export(1, "docx")
        assert result == template_binary
