"""DOCX surgical editing — replace text in specific paragraphs while preserving formatting.

Two-phase system:
1. Import time: build_paragraph_map() scans the DOCX and records which paragraph index
   corresponds to which section/company/bullet.
2. Generation time: apply_edits() loads the DOCX, replaces text at mapped indices,
   and serializes back to bytes. All fonts, spacing, bold, colors preserved.
"""

from __future__ import annotations

import io
import re
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn

from shared.scraping.resume_parser import SECTION_HEADERS, DATE_PATTERN


# ──────────────────────────────────────────────
#  Phase 1: Build paragraph map at import time
# ──────────────────────────────────────────────

def build_paragraph_map(doc: Document) -> dict:
    """Walk all paragraphs in a DOCX and build a structural map.

    Returns a dict with section locations and paragraph indices
    so that generation-time code knows exactly which paragraphs to edit.
    """
    paragraphs = doc.paragraphs
    total = len(paragraphs)

    current_section: str | None = None
    summary_indices: list[int] = []
    skills_indices: list[int] = []
    education_indices: list[int] = []
    projects_indices: list[int] = []
    roles: list[dict] = []
    current_role: dict | None = None

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Check if this is a section header
        is_bold = any(run.bold for run in para.runs) if para.runs else False
        style_name = para.style.name if para.style else ""
        is_heading = "heading" in (style_name or "").lower()
        matched_section = _match_section_header(text, is_bold or is_heading)

        if matched_section:
            # Save any in-progress role
            if current_role:
                roles.append(current_role)
                current_role = None
            current_section = matched_section
            continue

        # Accumulate paragraphs by section
        if current_section == "summary":
            summary_indices.append(i)

        elif current_section == "experience":
            if _is_role_header(text, para):
                if current_role:
                    roles.append(current_role)
                company, title = _extract_company_title(text)
                current_role = {
                    "company": company,
                    "title": title,
                    "header_index": i,
                    "bullet_indices": [],
                    "bullet_texts": [],
                }
            elif current_role:
                current_role["bullet_indices"].append(i)
                current_role["bullet_texts"].append(text)

        elif current_section == "skills":
            skills_indices.append(i)

        elif current_section == "education":
            education_indices.append(i)

        elif current_section == "projects":
            projects_indices.append(i)

    # Don't forget the last role
    if current_role:
        roles.append(current_role)

    return {
        "format": "docx",
        "total_paragraphs": total,
        "sections": {
            "summary": {"paragraph_indices": summary_indices},
            "experience": {"roles": roles},
            "skills": {"paragraph_indices": skills_indices},
            "education": {"paragraph_indices": education_indices},
            "projects": {"paragraph_indices": projects_indices},
        },
    }


def _match_section_header(text: str, is_formatted: bool) -> str | None:
    """Check if text matches a known section header."""
    if not is_formatted:
        return None
    for section_name, pattern in SECTION_HEADERS.items():
        if pattern.match(text):
            return section_name
    return None


def _is_role_header(text: str, para) -> bool:
    """Detect if a paragraph is a role header (Company | Title  Dates)."""
    # Must contain a separator (tab, pipe, em-dash) and a date
    has_separator = "\t" in text or " | " in text or " — " in text or " – " in text
    has_date = bool(DATE_PATTERN.search(text))
    is_not_bullet = not text.startswith("-") and not text.startswith("•")
    style_name = para.style.name if para.style else ""
    is_not_list = "list" not in (style_name or "").lower()
    return has_separator and has_date and is_not_bullet and is_not_list


def _extract_company_title(text: str) -> tuple[str, str]:
    """Extract company and title from a role header line."""
    # Split on tab, pipe, em-dash
    parts = re.split(r"\t|\s*[|–—]\s*", text, maxsplit=2)
    company = parts[0].strip() if parts else text
    title = parts[1].strip() if len(parts) > 1 else ""
    return company, title


# ──────────────────────────────────────────────
#  Phase 2: Apply edits at generation time
# ──────────────────────────────────────────────

def apply_edits(docx_bytes: bytes, paragraph_map: dict, edits: dict) -> bytes | None:
    """Load a DOCX, apply LLM edits to mapped paragraphs, return modified DOCX bytes.

    Returns None if the paragraph map is stale (count mismatch) — caller should
    fall back to text-based assembly.
    """
    doc = Document(io.BytesIO(docx_bytes))

    # Integrity check
    if len(doc.paragraphs) != paragraph_map.get("total_paragraphs"):
        return None

    sections = paragraph_map.get("sections", {})

    # 1. Replace summary
    new_summary = edits.get("summary")
    if new_summary:
        summary_indices = sections.get("summary", {}).get("paragraph_indices", [])
        for j, idx in enumerate(summary_indices):
            if j == 0:
                _replace_text(doc.paragraphs[idx], new_summary)
            else:
                # Clear extra summary paragraphs
                _replace_text(doc.paragraphs[idx], "")

    # 2. Replace experience bullets
    experience_edits = {
        e["company"].lower(): e for e in edits.get("experience_edits", [])
    }
    for role in sections.get("experience", {}).get("roles", []):
        role_company = (role.get("company") or "").lower()
        # Find matching edit by company name (case-insensitive, substring match)
        matched_edit = None
        for edit_company, edit in experience_edits.items():
            if edit_company in role_company or role_company in edit_company:
                matched_edit = edit
                break

        if not matched_edit:
            continue

        new_bullets = matched_edit.get("bullets", [])
        old_indices = role.get("bullet_indices", [])

        for j, idx in enumerate(old_indices):
            if j < len(new_bullets):
                bullet_text = new_bullets[j]
                # Strip leading "- " if present — the paragraph style handles the bullet
                if bullet_text.startswith("- "):
                    bullet_text = bullet_text[2:]
                elif bullet_text.startswith("• "):
                    bullet_text = bullet_text[2:]
                _replace_text(doc.paragraphs[idx], bullet_text)
            else:
                # Fewer new bullets than original — clear excess
                _replace_text(doc.paragraphs[idx], "")

        # If more new bullets than original — insert after last existing bullet
        if len(new_bullets) > len(old_indices) and old_indices:
            last_idx = old_indices[-1]
            style_source = doc.paragraphs[old_indices[0]]
            for extra_bullet in new_bullets[len(old_indices):]:
                if extra_bullet.startswith("- "):
                    extra_bullet = extra_bullet[2:]
                elif extra_bullet.startswith("• "):
                    extra_bullet = extra_bullet[2:]
                _insert_paragraph_after(doc, last_idx, extra_bullet, style_source)
                last_idx += 1  # approximate — inserted paragraphs shift indices

    # Serialize
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ──────────────────────────────────────────────
#  Text replacement with format preservation
# ──────────────────────────────────────────────

def _replace_text(paragraph, new_text: str) -> None:
    """Replace paragraph text while preserving the formatting of the first run."""
    if not paragraph.runs:
        paragraph.text = new_text
        return

    # Capture formatting XML from first run
    first_rpr = paragraph.runs[0]._element.find(qn("w:rPr"))
    saved_rpr = deepcopy(first_rpr) if first_rpr is not None else None

    # Clear all runs
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    # Add new run with preserved formatting
    new_run = paragraph.add_run(new_text)
    if saved_rpr is not None:
        existing_rpr = new_run._element.find(qn("w:rPr"))
        if existing_rpr is not None:
            new_run._element.replace(existing_rpr, saved_rpr)
        else:
            new_run._element.insert(0, saved_rpr)


def _insert_paragraph_after(doc: Document, after_index: int, text: str, style_source) -> None:
    """Insert a new paragraph after the given index, cloning style from source."""
    reference = doc.paragraphs[after_index]._element
    new_para = deepcopy(style_source._element)

    # Clear text and set new content
    for run_elem in new_para.findall(qn("w:r")):
        new_para.remove(run_elem)

    # Create a new run with the text
    from docx.oxml import OxmlElement

    run_elem = OxmlElement("w:r")
    # Copy run formatting from style source's first run
    source_rpr = style_source._element.find(qn("w:r"))
    if source_rpr is not None:
        source_run_rpr = source_rpr.find(qn("w:rPr"))
        if source_run_rpr is not None:
            run_elem.append(deepcopy(source_run_rpr))
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    text_elem.set(qn("xml:space"), "preserve")
    run_elem.append(text_elem)
    new_para.append(run_elem)

    # Insert after reference in the document body
    reference.addnext(new_para)
