"""DOCX exporter — converts markdown to Word document via python-docx."""

import io
import re

from docx import Document
from docx.shared import Pt


class DocxExporter:
    def export(self, content: str, metadata: dict) -> bytes:
        doc = Document()
        _populate_document(doc, content)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def format_name(self) -> str:
        return "docx"


def _populate_document(doc: Document, markdown_text: str) -> None:
    """Parse markdown and add content to the DOCX document."""
    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("**") and stripped.endswith("**"):
            # Bold-only lines (e.g., "**Languages:** Python")
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            para = doc.add_paragraph()
            run = para.add_run(clean)
            run.bold = True
        else:
            # Strip inline markdown formatting
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            doc.add_paragraph(clean)
